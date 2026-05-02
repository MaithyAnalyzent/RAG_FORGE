from __future__ import annotations

from pathlib import Path
from typing import List

from src.interfaces.document_loader import Document


class DocxLoader:
    """Extracts text and table content from Microsoft Word (.docx) files."""

    supported_extensions = [".docx"]

    def load(self, path: Path) -> List[Document]:
        from docx import Document as WordDoc

        doc = WordDoc(str(path))
        parts: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return [
            Document(
                content="\n\n".join(parts),
                source=path.name,
                metadata={"path": str(path), "format": "docx"},
            )
        ]
